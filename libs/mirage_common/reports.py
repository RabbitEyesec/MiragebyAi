"""Versioned, provenance-preserving Mirage case reports and signed packages."""
from __future__ import annotations

import hashlib
import io
import json
import re
import zipfile
from dataclasses import asdict, dataclass
from datetime import UTC, datetime
from pathlib import Path
from typing import Any, Literal
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    KeepTogether,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)

from mirage_common.evidence_export import (
    FIXED_ZIP_TIME,
    MANIFEST_VERSION,
    SIGNING_ALGORITHM,
    ManifestSigner,
    sign_manifest,
    verify_signature,
)
from mirage_common.trust_anchor import resolve_trusted_key
from mirage_contracts.envelope import canonical_json_bytes

REPORT_TEMPLATE_VERSION = "mirage-case-report/1.0"
REPORT_SCHEMA_VERSION = "mirage.case-report/1.0"
REPORT_GENERATOR_VERSION = "0.1.0"
CLASSIFICATIONS = (
    "OBSERVED_FACT",
    "DETERMINISTIC_CORRELATION",
    "AI_INFERENCE",
    "ANALYST_NOTE",
)
REPORT_SECTIONS = (
    "Cover page",
    "Case identity",
    "Executive summary",
    "Scope and controlled-lab boundary",
    "Case lifecycle",
    "Session summary",
    "Detection and correlation",
    "Steering decision",
    "Sandbox identity and fingerprint result",
    "Investigation timeline",
    "Behaviour profile",
    "Skill assessment",
    "Observed commands and tooling",
    "Network indicators",
    "Artifact activity",
    "Canary activity",
    "AI proposals",
    "Policy decisions",
    "Sandbox actions",
    "Analyst directives",
    "Analyst direct messages",
    "Evidence inventory",
    "Evidence verification summary",
    "Collection gaps",
    "Limitations",
    "Integrity and signing information",
    "Independent verification instructions",
)
_SECRET = re.compile(
    r"(?i)\b(password|passphrase|api[ _-]?key|access[ _-]?token|refresh[ _-]?token|"
    r"token|secret|credential)\b(\s*[:=]\s*)([^\s,;]+)"
)


@dataclass(frozen=True)
class Provenance:
    event_ids: tuple[str, ...] = ()
    evidence_ids: tuple[str, ...] = ()
    timeline_ids: tuple[str, ...] = ()
    source_identity: str | None = None
    source_sequence: int | None = None


@dataclass(frozen=True)
class ReportStatement:
    category: Literal[
        "OBSERVED_FACT",
        "DETERMINISTIC_CORRELATION",
        "AI_INFERENCE",
        "ANALYST_NOTE",
    ]
    text: str
    provenance: Provenance
    confidence: float | None = None
    contradictory_event_ids: tuple[str, ...] = ()
    uncertainty: str | None = None
    model: str | None = None
    provider: str | None = None


@dataclass(frozen=True)
class ReportSection:
    title: str
    statements: tuple[ReportStatement, ...]


@dataclass(frozen=True)
class ReportModel:
    metadata: dict[str, Any]
    sections: tuple[ReportSection, ...]


@dataclass(frozen=True)
class ReportArtifacts:
    model: ReportModel
    pdf: bytes
    docx: bytes
    case_json: bytes


@dataclass(frozen=True)
class ReportPackage:
    package: bytes
    manifest_sha256: str
    package_sha256: str


@dataclass(frozen=True)
class ReportVerification:
    valid: bool
    errors: tuple[str, ...]
    manifest_sha256: str | None


def redact_text(value: Any, *, maximum: int = 4096) -> str:
    """Return bounded display text without credential-shaped material."""
    if value is None:
        return "Not recorded"
    if isinstance(value, dict | list | tuple):
        text = json.dumps(value, sort_keys=True, ensure_ascii=False, default=str)
    else:
        text = str(value)
    text = text.replace("\x00", "\N{REPLACEMENT CHARACTER}")
    text = _SECRET.sub(lambda match: f"{match.group(1)}{match.group(2)}[REDACTED]", text)
    return text[:maximum] + ("…" if len(text) > maximum else "")


def _iso(value: Any) -> str:
    if isinstance(value, datetime):
        return value.astimezone(UTC).isoformat().replace("+00:00", "Z")
    return redact_text(value)


def _reference(item: dict[str, Any], default_source: str) -> Provenance:
    event_ids = item.get("event_ids") or item.get("source_event_ids")
    if not event_ids:
        event_ids = item.get("related_event_ids") or item.get("supporting_event_ids") or []
    evidence_ids = item.get("evidence_ids") or item.get("evidence_references") or []
    evidence_id = item.get("evidence_id")
    if evidence_id:
        evidence_ids = [*evidence_ids, evidence_id]
    timeline_id = item.get("timeline_id") or item.get("item_id")
    sequence = item.get("source_sequence")
    return Provenance(
        event_ids=tuple(sorted(str(value) for value in event_ids)),
        evidence_ids=tuple(sorted(str(value) for value in evidence_ids)),
        timeline_ids=(str(timeline_id),) if timeline_id else (),
        source_identity=str(item.get("source_id") or default_source),
        source_sequence=int(sequence) if sequence is not None else None,
    )


def _fact(text: str, item: dict[str, Any], source: str) -> ReportStatement:
    return ReportStatement("OBSERVED_FACT", redact_text(text), _reference(item, source))


def _correlation(text: str, item: dict[str, Any], source: str) -> ReportStatement:
    return ReportStatement(
        "DETERMINISTIC_CORRELATION", redact_text(text), _reference(item, source)
    )


def _empty(title: str) -> ReportStatement:
    return _correlation(
        f"No {title.lower()} records were present in the selected source projection.",
        {},
        "dashboard-read-model",
    )


def _rows(
    title: str,
    values: list[dict[str, Any]],
    *,
    source: str,
    formatter: Any,
    category: str = "OBSERVED_FACT",
) -> tuple[ReportStatement, ...]:
    if not values:
        return (_empty(title),)
    ordered = sorted(
        values,
        key=lambda item: (
            str(item.get("created_at") or item.get("event_time") or item.get("at") or ""),
            str(
                item.get("id")
                or item.get("item_id")
                or item.get("evidence_id")
                or item.get("action_id")
                or ""
            ),
        ),
    )
    make = _correlation if category == "DETERMINISTIC_CORRELATION" else _fact
    return tuple(make(formatter(item), item, source) for item in ordered)


def build_report_model(
    data: dict[str, Any],
    *,
    report_id: str,
    export_id: str,
    created_by: str,
    created_at: str,
    build_hash: str,
    source_projection_version: int,
    evidence_manifest_id: str,
    timezone_name: str = "UTC",
    limitations: list[str] | None = None,
) -> ReportModel:
    """Build all 27 sections with a classification on every statement."""
    case = dict(data.get("case") or {})
    sessions = [dict(item) for item in data.get("sessions") or []]
    timeline = [dict(item) for item in data.get("timeline") or []]
    evidence = [dict(item) for item in data.get("evidence") or []]
    gaps = [dict(item) for item in data.get("gaps") or []]
    directives = [dict(item) for item in data.get("directives") or []]
    messages = [dict(item) for item in data.get("messages") or []]
    proposals = [dict(item) for item in data.get("ai") or []]
    decisions = [dict(item) for item in data.get("policy") or []]
    graph = [dict(item) for item in data.get("graph_nodes") or []]
    detections = [
        item for item in graph if item.get("node_type") in {"DETECTION", "CORRELATION"}
    ]
    sandboxes = [dict(item) for item in data.get("sandboxes") or []]
    actions = [dict(item) for item in data.get("sandbox_actions") or []]
    behaviours = [dict(item) for item in data.get("behaviour") or []]
    artifacts = [dict(item) for item in data.get("artifacts") or []]
    canaries = [dict(item) for item in data.get("canaries") or []]
    network = [dict(item) for item in data.get("network_indicators") or []]
    commands = [dict(item) for item in data.get("observed_commands") or []]

    case_id = redact_text(case.get("case_id") or "unknown")
    case_source = {"source_id": "cases", "source_sequence": case.get("version")}
    metadata = {
        "template_version": REPORT_TEMPLATE_VERSION,
        "report_schema_version": REPORT_SCHEMA_VERSION,
        "generator_version": REPORT_GENERATOR_VERSION,
        "build_hash": build_hash,
        "report_id": report_id,
        "creation_actor": created_by,
        "creation_time": created_at,
        "source_projection_version": source_projection_version,
        "evidence_manifest_id": evidence_manifest_id,
        "export_id": export_id,
        "timezone": timezone_name,
    }
    verified = sum(item.get("verification_status") == "VERIFIED" for item in evidence)
    sections: dict[str, tuple[ReportStatement, ...]] = {
        "Cover page": (
            _fact(
                f"Mirage controlled-lab case report for case {case_id}.",
                case_source,
                "cases",
            ),
        ),
        "Case identity": (
            _fact(
                f"Case {case_id}; state {case.get('state', 'UNKNOWN')}; "
                f"severity {case.get('severity', 'UNKNOWN')}; owner "
                f"{case.get('owner') or 'unassigned'}; version {case.get('version', 0)}.",
                case_source,
                "cases",
            ),
        ),
        "Executive summary": (
            _correlation(
                f"The selected projection contains {len(sessions)} session(s), "
                f"{len(timeline)} timeline transition(s), and {len(evidence)} evidence object(s).",
                case_source,
                "dashboard-read-model",
            ),
        ),
        "Scope and controlled-lab boundary": (
            _correlation(
                "This report describes only activity recorded inside the authorised Mirage "
                "controlled-lab boundary; absence of a record is not proof of absence.",
                {},
                "report-template",
            ),
        ),
        "Case lifecycle": _rows(
            "Case lifecycle",
            timeline,
            source="case-state-transitions",
            formatter=lambda item: (
                f"{_iso(item.get('at'))}: {item.get('from_state') or 'NONE'} → "
                f"{item.get('to_state') or 'UNKNOWN'}; actor {item.get('actor') or 'unknown'}; "
                f"reason {item.get('reason') or 'not recorded'}."
            ),
        ),
        "Session summary": _rows(
            "Session summary",
            sessions,
            source="sessions",
            formatter=lambda item: (
                f"Session {item.get('session_id')}: protocol {item.get('protocol')}; "
                f"status {item.get('status')}; created {_iso(item.get('created_at'))}."
            ),
        ),
        "Detection and correlation": _rows(
            "Detection and correlation",
            detections,
            source="dashboard-graph",
            formatter=lambda item: f"{item.get('node_type')}: {item.get('label')}",
            category="DETERMINISTIC_CORRELATION",
        ),
        "Steering decision": _rows(
            "Steering decision",
            [dict(item) for item in data.get("routing_decisions") or []],
            source="routing-decisions",
            formatter=lambda item: (
                f"Decision {item.get('decision_id') or item.get('id')}: "
                f"{item.get('protocol')} traffic target {item.get('target')}."
            ),
            category="DETERMINISTIC_CORRELATION",
        ),
        "Sandbox identity and fingerprint result": _rows(
            "Sandbox identity and fingerprint result",
            sandboxes,
            source="sandbox-instances",
            formatter=lambda item: (
                f"Sandbox {item.get('sandbox_id')}: image {item.get('image_id')}; "
                f"state {item.get('status') or item.get('state')}; fingerprint "
                f"{item.get('fingerprint_status') or 'not recorded'}."
            ),
        ),
        "Investigation timeline": _rows(
            "Investigation timeline",
            [dict(item) for item in data.get("dashboard_timeline") or timeline],
            source="dashboard-timeline",
            formatter=lambda item: (
                f"{_iso(item.get('event_time') or item.get('at'))}: "
                f"{item.get('label') or item.get('to_state') or 'timeline item'} — "
                f"{item.get('description') or item.get('reason') or ''}"
            ),
        ),
        "Behaviour profile": _rows(
            "Behaviour profile",
            behaviours,
            source="behaviour-profile",
            formatter=lambda item: (
                f"{item.get('dimension') or item.get('name')}: "
                f"{item.get('value') or item.get('label')}; confidence "
                f"{item.get('confidence') if item.get('confidence') is not None else 'not recorded'}."
            ),
            category="DETERMINISTIC_CORRELATION",
        ),
        "Skill assessment": _rows(
            "Skill assessment",
            [item for item in behaviours if item.get("dimension") == "skill"],
            source="behaviour-profile",
            formatter=lambda item: (
                f"Skill assessment {item.get('value') or item.get('label')}; "
                f"basis {item.get('basis') or 'recorded behaviour features'}."
            ),
            category="DETERMINISTIC_CORRELATION",
        ),
        "Observed commands and tooling": _rows(
            "Observed commands and tooling",
            commands,
            source="spider-observation",
            formatter=lambda item: (
                f"Observed command/tool record: {item.get('display') or item.get('command') or item.get('label')}."
            ),
        ),
        "Network indicators": _rows(
            "Network indicators",
            network,
            source="network-observation",
            formatter=lambda item: (
                f"Network indicator {item.get('type') or 'unknown'}: "
                f"{item.get('value') or item.get('label')}."
            ),
        ),
        "Artifact activity": _rows(
            "Artifact activity",
            artifacts,
            source="artifact-ledger",
            formatter=lambda item: (
                f"Artifact {item.get('artifact_id')}: scan {item.get('scan_status')}; "
                f"deployment {item.get('deployment_status') or 'not deployed'}."
            ),
        ),
        "Canary activity": _rows(
            "Canary activity",
            canaries,
            source="canary-callbacks",
            formatter=lambda item: (
                f"Canary {item.get('token_id') or item.get('canary_id')}: "
                f"classification {item.get('classification')}; callback "
                f"{_iso(item.get('callback_time') or item.get('created_at'))}."
            ),
        ),
        "Policy decisions": _rows(
            "Policy decisions",
            decisions,
            source="policy-decisions",
            formatter=lambda item: (
                f"Policy decision {item.get('decision_id')}: {item.get('decision')}; "
                f"reason codes {item.get('reason_codes') or []}."
            ),
            category="DETERMINISTIC_CORRELATION",
        ),
        "Sandbox actions": _rows(
            "Sandbox actions",
            actions,
            source="sandbox-actions",
            formatter=lambda item: (
                f"Action {item.get('action_id')}: {item.get('action_type')}; "
                f"status {item.get('status')}."
            ),
        ),
        "Analyst directives": _rows(
            "Analyst directives",
            directives,
            source="analyst-directives",
            formatter=lambda item: (
                f"Directive {item.get('directive_id') or item.get('id')}: "
                f"{item.get('objective')}; status {item.get('status')}."
            ),
        ),
        "Analyst direct messages": _rows(
            "Analyst direct messages",
            messages,
            source="analyst-messages",
            formatter=lambda item: (
                f"Message {item.get('message_id') or item.get('id')} to "
                f"{item.get('surface')}; status {item.get('status')}; output tag "
                f"{item.get('output_tag')}."
            ),
        ),
        "Evidence inventory": _rows(
            "Evidence inventory",
            evidence,
            source="evidence-ledger",
            formatter=lambda item: (
                f"Evidence {item.get('evidence_id')}: type {item.get('evidence_type')}; "
                f"SHA-256 {item.get('sha256')}; verification "
                f"{item.get('verification_status')}."
            ),
        ),
        "Evidence verification summary": (
            _correlation(
                f"{verified} of {len(evidence)} evidence object(s) are VERIFIED in the ledger.",
                {},
                "evidence-ledger",
            ),
        ),
        "Collection gaps": _rows(
            "Collection gaps",
            gaps,
            source="evidence-collection-gaps",
            formatter=lambda item: (
                f"Gap {item.get('gap_id')}: {item.get('reason')}; required "
                f"{item.get('required')}; resolution {item.get('resolution') or 'unresolved'}."
            ),
        ),
        "Limitations": tuple(
            _correlation(text, {}, "report-generation-policy")
            for text in (
                limitations
                or [
                    "This export is evidence-ready and is not itself a claim of court admissibility.",
                    "Timestamps and signatures inherit the trust level stated by the evidence manifest.",
                ]
            )
        ),
        "Integrity and signing information": (
            _correlation(
                f"Report {report_id} references export {export_id}, evidence manifest "
                f"{evidence_manifest_id}, build {build_hash}, and projection "
                f"version {source_projection_version}.",
                {},
                "report-generator",
            ),
        ),
        "Independent verification instructions": (
            _correlation(
                "Run scripts/verify-report-package against the package and a separately "
                "obtained public key; inspect every reported mismatch before relying on it.",
                {},
                "report-template",
            ),
        ),
    }

    inference: list[ReportStatement] = []
    for item in sorted(proposals, key=lambda value: str(value.get("proposal_id", ""))):
        inference.append(
            ReportStatement(
                "AI_INFERENCE",
                redact_text(
                    f"Proposal {item.get('proposal_id')}: {item.get('rationale')}; "
                    f"proposed action {item.get('action_type')}."
                ),
                _reference(item, "ai-proposals"),
                confidence=float(item.get("confidence") or 0),
                contradictory_event_ids=tuple(
                    sorted(str(value) for value in item.get("contradictory_event_ids") or [])
                ),
                uncertainty=redact_text(
                    item.get("uncertainty")
                    or "No explicit uncertainty statement was retained by the provider."
                ),
                model=redact_text(item.get("model") or data.get("ai_model") or "not disclosed"),
                provider=redact_text(
                    item.get("provider") or data.get("ai_provider") or "not disclosed"
                ),
            )
        )
    sections["AI proposals"] = tuple(inference) or (_empty("AI proposals"),)
    return ReportModel(
        metadata=metadata,
        sections=tuple(ReportSection(title, sections[title]) for title in REPORT_SECTIONS),
    )


def report_model_json(model: ReportModel) -> bytes:
    value = {
        "metadata": model.metadata,
        "sections": [
            {
                "title": section.title,
                "statements": [
                    {
                        **asdict(statement),
                        "provenance": asdict(statement.provenance),
                    }
                    for statement in section.statements
                ],
            }
            for section in model.sections
        ],
    }
    return canonical_json_bytes(value)


def _provenance_text(statement: ReportStatement) -> str:
    parts: list[str] = []
    provenance = statement.provenance
    if provenance.event_ids:
        parts.append("events=" + ",".join(provenance.event_ids))
    if provenance.evidence_ids:
        parts.append("evidence=" + ",".join(provenance.evidence_ids))
    if provenance.timeline_ids:
        parts.append("timeline=" + ",".join(provenance.timeline_ids))
    if provenance.source_identity:
        parts.append(f"source={provenance.source_identity}")
    if provenance.source_sequence is not None:
        parts.append(f"sequence={provenance.source_sequence}")
    if statement.category == "AI_INFERENCE":
        parts.extend(
            (
                f"confidence={statement.confidence}",
                "contradictory-events="
                + (",".join(statement.contradictory_event_ids) or "none recorded"),
                f"uncertainty={statement.uncertainty}",
                f"provider={statement.provider}",
                f"model={statement.model}",
            )
        )
    return " · ".join(parts) or "source=report-generation-policy"


def render_pdf(model: ReportModel) -> bytes:
    buffer = io.BytesIO()
    document = BaseDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title=f"Mirage report {model.metadata['report_id']}",
        author="Mirage report worker",
    )
    frame = Frame(document.leftMargin, document.bottomMargin, document.width, document.height)

    def footer(canvas: Any, doc: Any) -> None:
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#53616f"))
        canvas.drawString(18 * mm, 10 * mm, "Mirage controlled-lab report · UTC")
        canvas.drawRightString(A4[0] - 18 * mm, 10 * mm, f"Page {doc.page}")
        canvas.restoreState()

    document.addPageTemplates([PageTemplate(id="report", frames=[frame], onPage=footer)])
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            "Cover",
            parent=styles["Title"],
            alignment=TA_CENTER,
            fontName="Helvetica-Bold",
            fontSize=24,
            leading=30,
            textColor=colors.HexColor("#10242b"),
            spaceAfter=18,
        )
    )
    styles.add(
        ParagraphStyle(
            "Statement",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=9,
            leading=13,
            splitLongWords=True,
        )
    )
    styles.add(
        ParagraphStyle(
            "Provenance",
            parent=styles["BodyText"],
            fontName="Courier",
            fontSize=6.5,
            leading=9,
            textColor=colors.HexColor("#4f5f68"),
            splitLongWords=True,
        )
    )
    category_colours = {
        "OBSERVED_FACT": "#d9efe8",
        "DETERMINISTIC_CORRELATION": "#dce8f6",
        "AI_INFERENCE": "#f5dfef",
        "ANALYST_NOTE": "#f6e9c7",
    }
    story: list[Any] = [
        Spacer(1, 28 * mm),
        Paragraph("MIRAGE CASE REPORT", styles["Cover"]),
        Paragraph(
            escape(redact_text(model.metadata["report_id"])),
            styles["Heading2"],
        ),
        Spacer(1, 14 * mm),
        Paragraph(
            "Classified statements · evidence-linked · independently verifiable",
            styles["BodyText"],
        ),
        PageBreak(),
        Paragraph("Table of contents", styles["Heading1"]),
    ]
    for number, section in enumerate(model.sections, 1):
        story.append(Paragraph(f"{number}. {escape(section.title)}", styles["BodyText"]))
    story.append(PageBreak())
    for number, section in enumerate(model.sections, 1):
        story.append(Paragraph(f"{number}. {escape(section.title)}", styles["Heading1"]))
        for statement in section.statements:
            label = Paragraph(f"<b>{statement.category}</b>", styles["Statement"])
            body = Paragraph(escape(statement.text), styles["Statement"])
            provenance = Paragraph(escape(_provenance_text(statement)), styles["Provenance"])
            table = Table(
                [[label], [body], [provenance]],
                colWidths=[document.width],
                repeatRows=1,
                splitByRow=True,
            )
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(category_colours[statement.category])),
                        ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#9aa8ae")),
                        ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#d5dcdf")),
                        ("LEFTPADDING", (0, 0), (-1, -1), 7),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                        ("TOPPADDING", (0, 0), (-1, -1), 5),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ]
                )
            )
            story.extend((KeepTogether([table]), Spacer(1, 4 * mm)))
    document.build(story)
    return buffer.getvalue()


def _add_docx_page_number(paragraph: Any) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, end))


def render_docx(model: ReportModel) -> bytes:
    document = Document()
    document.core_properties.title = f"Mirage report {model.metadata['report_id']}"
    document.core_properties.author = "Mirage report worker"
    document.core_properties.subject = "Controlled-lab case report"
    document.core_properties.comments = "No remote resources are embedded."
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    footer = section.footer.paragraphs[0]
    footer.add_run("Mirage controlled-lab report · UTC · Page ")
    _add_docx_page_number(footer)

    styles = document.styles
    for name, colour in (
        ("OBSERVED_FACT", RGBColor(0x19, 0x66, 0x4A)),
        ("DETERMINISTIC_CORRELATION", RGBColor(0x24, 0x5A, 0x89)),
        ("AI_INFERENCE", RGBColor(0x8A, 0x34, 0x73)),
        ("ANALYST_NOTE", RGBColor(0x8A, 0x5A, 0x10)),
    ):
        if name not in styles:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = "Arial"
            style.font.size = Pt(10)
            style.font.color.rgb = colour

    document.add_heading("MIRAGE CASE REPORT", level=0)
    document.add_paragraph(redact_text(model.metadata["report_id"]), style="Subtitle")
    document.add_paragraph(
        "Classified statements · evidence-linked · independently verifiable"
    )
    document.add_page_break()
    document.add_heading("Table of contents", level=1)
    toc = document.add_paragraph()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate.text = "Right-click and update field to populate the table of contents."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    toc._p.extend((begin, instruction, separate, end))
    document.add_page_break()
    for number, report_section in enumerate(model.sections, 1):
        document.add_heading(f"{number}. {report_section.title}", level=1)
        for statement in report_section.statements:
            paragraph = document.add_paragraph(style=statement.category)
            label = paragraph.add_run(f"{statement.category}\n")
            label.bold = True
            paragraph.add_run(statement.text)
            provenance = document.add_paragraph(_provenance_text(statement))
            provenance.style = document.styles["Caption"]
            for run in provenance.runs:
                run.font.name = "Consolas"
                run.font.size = Pt(7)
    final_section = document.add_section(WD_SECTION.CONTINUOUS)
    final_section.header.is_linked_to_previous = True
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def create_report_artifacts(model: ReportModel) -> ReportArtifacts:
    case_json = report_model_json(model)
    return ReportArtifacts(model, render_pdf(model), render_docx(model), case_json)


def _write_member(zf: zipfile.ZipFile, name: str, data: bytes) -> None:
    info = zipfile.ZipInfo(name, FIXED_ZIP_TIME)
    info.compress_type = zipfile.ZIP_DEFLATED
    info.external_attr = 0o100644 << 16
    zf.writestr(info, data)


def create_report_package(
    artifacts: ReportArtifacts,
    *,
    signer: ManifestSigner,
    export_mode: str,
    evidence_manifest: dict[str, Any],
    evidence_objects: dict[str, bytes] | None = None,
) -> ReportPackage:
    if export_mode not in {"METADATA_ONLY", "SELECTED_EVIDENCE", "COMPLETE_CASE"}:
        raise ValueError("unsupported report export mode")
    objects = evidence_objects or {}
    evidence = sorted(
        [dict(item) for item in evidence_manifest.get("evidence_objects") or []],
        key=lambda item: str(item.get("evidence_id", "")),
    )
    expected_evidence = {str(item["evidence_id"]) for item in evidence}
    if export_mode == "METADATA_ONLY" and objects:
        raise ValueError("METADATA_ONLY report cannot embed evidence bytes")
    if export_mode == "COMPLETE_CASE" and set(objects) != expected_evidence:
        raise ValueError("COMPLETE_CASE report must include every manifest evidence object")
    if not set(objects).issubset(expected_evidence):
        raise ValueError("report contains evidence absent from the evidence manifest")

    verification_json = canonical_json_bytes(
        {
            "status": "GENERATED_PENDING_INDEPENDENT_VERIFICATION",
            "verifier": "mirage-report-worker",
            "report_id": artifacts.model.metadata["report_id"],
            "checks": [
                "canonical manifest",
                "report file hashes",
                "evidence manifest linkage",
                "asymmetric signature",
            ],
        }
    )
    verification_text = (
        b"Mirage Report Package Verification\n"
        b"Status: GENERATED_PENDING_INDEPENDENT_VERIFICATION\n"
        b"Run: scripts/verify-report-package <package.zip> [--public-key key.pem]\n"
    )
    readme = (
        b"MIRAGE REPORT PACKAGE\n\n"
        b"Verify this package independently before relying on it.\n"
        b"Do not treat AI_INFERENCE as observed fact.\n"
        b"No remote resource is required to open these report files.\n"
    )
    references = canonical_json_bytes(
        {
            "export_mode": export_mode,
            "included_evidence_ids": sorted(objects),
            "referenced_evidence_ids": sorted(expected_evidence),
        }
    )
    files = {
        "report.pdf": artifacts.pdf,
        "report.docx": artifacts.docx,
        "case.json": artifacts.case_json,
        "verification-report.json": verification_json,
        "verification-report.txt": verification_text,
        "README-verification.txt": readme,
        "evidence-references.json": references,
    }
    file_manifest = {
        name: {"sha256": hashlib.sha256(value).hexdigest(), "size_bytes": len(value)}
        for name, value in sorted(files.items())
    }
    manifest = {
        "manifest_version": MANIFEST_VERSION,
        "manifest_type": "MIRAGE_CASE_REPORT",
        "report": artifacts.model.metadata,
        "export_mode": export_mode,
        "evidence_manifest_id": artifacts.model.metadata["evidence_manifest_id"],
        "evidence_manifest_sha256": hashlib.sha256(
            canonical_json_bytes(evidence_manifest)
        ).hexdigest(),
        "evidence_objects": evidence,
        "report_files": file_manifest,
        "embedded_evidence_ids": sorted(objects),
        "signing": {
            "kms_key_arn": signer.key_id,
            "algorithm": SIGNING_ALGORITHM,
            "payload": "SHA256(canonical-report-manifest-without-signature)",
        },
    }
    signed = sign_manifest(manifest, signer)
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w") as zf:
        _write_member(zf, "manifest.json", signed.canonical_bytes)
        _write_member(zf, "manifest.sha256", (signed.manifest_sha256 + "\n").encode())
        _write_member(zf, "manifest.sig", signed.signature)
        _write_member(zf, "public-key.pem", signed.public_key_pem)
        for name, value in sorted(files.items()):
            _write_member(zf, name, value)
        for evidence_id, value in sorted(objects.items()):
            _write_member(zf, f"evidence/{evidence_id}", value)
    package = buffer.getvalue()
    return ReportPackage(
        package=package,
        manifest_sha256=signed.manifest_sha256,
        package_sha256=hashlib.sha256(package).hexdigest(),
    )


def verify_report_package(
    package: bytes,
    *,
    public_key_pem: bytes | None = None,
    trust_store_dir: Path | None = None,
    expected_evidence_manifest_sha256: str | None = None,
) -> ReportVerification:
    errors: list[str] = []
    manifest_hash: str | None = None
    try:
        with zipfile.ZipFile(io.BytesIO(package), "r") as zf:
            listed = zf.namelist()
            names = set(listed)
            if len(listed) != len(names):
                errors.append("package contains duplicate member names")
            required = {
                "manifest.json",
                "manifest.sha256",
                "manifest.sig",
                "public-key.pem",
                "report.pdf",
                "report.docx",
                "case.json",
                "verification-report.json",
                "verification-report.txt",
                "README-verification.txt",
                "evidence-references.json",
            }
            errors.extend(
                f"missing required package member: {name}" for name in sorted(required - names)
            )
            if "manifest.json" not in names:
                return ReportVerification(False, tuple(errors), None)
            manifest_bytes = zf.read("manifest.json")
            manifest_hash = hashlib.sha256(manifest_bytes).hexdigest()
            try:
                manifest = json.loads(manifest_bytes)
            except json.JSONDecodeError as exc:
                return ReportVerification(False, (f"manifest JSON invalid: {exc}",), manifest_hash)
            if canonical_json_bytes(manifest) != manifest_bytes:
                errors.append("manifest is not canonical JSON")
            if (
                "manifest.sha256" in names
                and zf.read("manifest.sha256").decode().strip() != manifest_hash
            ):
                errors.append("manifest hash mismatch")
            if manifest.get("manifest_type") != "MIRAGE_CASE_REPORT":
                errors.append("unexpected manifest type")
            if manifest.get("signing", {}).get("algorithm") != SIGNING_ALGORITHM:
                errors.append("unsupported signing algorithm")
            if (
                expected_evidence_manifest_sha256 is not None
                and manifest.get("evidence_manifest_sha256")
                != expected_evidence_manifest_sha256
            ):
                errors.append("evidence-manifest mismatch")
            declared_files = manifest.get("report_files")
            if not isinstance(declared_files, dict):
                errors.append("report_files must be an object")
                declared_files = {}
            expected = {
                "manifest.json",
                "manifest.sha256",
                "manifest.sig",
                "public-key.pem",
                *declared_files.keys(),
                *(f"evidence/{value}" for value in manifest.get("embedded_evidence_ids") or []),
            }
            errors.extend(f"unexpected file: {name}" for name in sorted(names - expected))
            for name, detail in sorted(declared_files.items()):
                if name not in names:
                    errors.append(f"missing file: {name}")
                    continue
                if not isinstance(detail, dict):
                    errors.append(f"invalid file manifest: {name}")
                    continue
                content = zf.read(name)
                if hashlib.sha256(content).hexdigest() != detail.get("sha256"):
                    errors.append(f"hash mismatch: {name}")
                if len(content) != detail.get("size_bytes"):
                    errors.append(f"size mismatch: {name}")
            embedded_key = zf.read("public-key.pem") if "public-key.pem" in names else None
            if "manifest.sig" not in names:
                errors.append("signature or verification key missing")
                trust = None
            else:
                trust = resolve_trusted_key(
                    explicit_key_bytes=public_key_pem,
                    trust_store_dir=trust_store_dir,
                    embedded_key_bytes=embedded_key,
                )
                errors.extend(trust.errors)
            if trust is not None and trust.key_bytes is not None:
                try:
                    verify_signature(
                        public_key_pem=trust.key_bytes,
                        digest=bytes.fromhex(manifest_hash),
                        signature=zf.read("manifest.sig"),
                    )
                except Exception as exc:  # noqa: BLE001
                    errors.append(f"signature verification failed: {type(exc).__name__}")
    except zipfile.BadZipFile:
        return ReportVerification(False, ("invalid report ZIP package",), None)
    return ReportVerification(not errors, tuple(errors), manifest_hash)
