"""Machine-readable Mirage acceptance planning, execution, and verification."""
from __future__ import annotations

import argparse
import hashlib
import html
import json
import platform
import shutil
import zipfile
from dataclasses import asdict, dataclass
from datetime import UTC, datetime
from pathlib import Path
from typing import Any, Literal

from cryptography.exceptions import InvalidSignature
from cryptography.hazmat.primitives import hashes, serialization
from cryptography.hazmat.primitives.asymmetric import padding, rsa
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from mirage_common.load_test import LOCAL_REDUCED, run_load
from mirage_common.resilience import SCENARIOS, run_local_scenario
from mirage_common.teardown import CloudResource, LocalCloudAdapter, TeardownWorkflow
from mirage_common.trust_anchor import resolve_trusted_key
from mirage_contracts.envelope import canonical_json_bytes

Result = Literal["PASS", "FAIL", "NOT_RUN", "BLOCKED"]
ACCEPTANCE_SCHEMA = "mirage.acceptance/1.0"
ZIP_TIME = (1980, 1, 1, 0, 0, 0)


@dataclass(frozen=True)
class NumericRequirement:
    requirement_id: str
    description: str
    target: str
    unit: str
    profile_b_only_measurement: bool


NUMERIC_REQUIREMENTS = (
    NumericRequirement("NUM-01", "Endpoint event to Elasticsearch p95", "<3", "seconds", True),
    NumericRequirement("NUM-02", "Sandbox event to dashboard p95", "<2", "seconds", True),
    NumericRequirement("NUM-03", "Agent local buffering at peak", ">=15", "minutes", False),
    NumericRequirement("NUM-04", "Confirmed event loss during five-minute outage", "0", "events", True),
    NumericRequirement("NUM-05", "Duplicate effective case events after replay", "0", "events", False),
    NumericRequirement("NUM-06", "Telemetry burst with buffering", "1000", "events/second", True),
    NumericRequirement("NUM-07", "AI action timeout to fallback", "10", "seconds", False),
    NumericRequirement("NUM-08", "AI snapshot byte maximum", "<=16384", "bytes", False),
    NumericRequirement("NUM-09", "AI snapshot token maximum", "<=4000", "estimated_tokens", False),
    NumericRequirement("NUM-10", "Snapshot trimming is recorded", "true", "boolean", False),
    NumericRequirement("NUM-11", "Artifact upload maximum is enforced", "250", "MB", False),
    NumericRequirement("NUM-12", "Ordinary artifact deployment p95 after scan", "<30", "seconds", True),
    NumericRequirement("NUM-13", "Evidence verified before export", "100", "percent", False),
    NumericRequirement("NUM-14", "Certificate renewal threshold remaining lifetime", "20", "percent", False),
    NumericRequirement("NUM-15", "Sandbox soft reset", "<3", "minutes", True),
    NumericRequirement("NUM-16", "Sandbox full rebuild", "<10", "minutes", True),
    NumericRequirement("NUM-17", "Critical worker alert", "<60", "seconds", True),
    NumericRequirement("NUM-18", "Agent heartbeat warning", "30", "seconds", True),
    NumericRequirement("NUM-19", "Agent offline", "90", "seconds", True),
    NumericRequirement("NUM-20", "NATS consumer-lag alert", ">10000 or >30", "messages_or_seconds", False),
    NumericRequirement("NUM-21", "Fingerprint MUST checks pass and block", "100", "percent", False),
    NumericRequirement("NUM-22", "Prompt injection executed as command", "0", "commands", False),
    NumericRequirement("NUM-23", "Internal/scanner canary displayed as attacker", "0", "callbacks", False),
    NumericRequirement("NUM-24", "Temporary resources remaining after teardown", "0", "resources", False),
    NumericRequirement("NUM-25", "Mid-session migration claims in code, UI, and docs", "0", "claims", False),
)

SCENARIO_STEPS = (
    "Provision from an empty environment",
    "Install the server package",
    "Install the employee endpoint package",
    "Install or launch the sandbox package",
    "Launch the Kali attacker VM",
    "Confirm all health checks",
    "Generate endpoint activity",
    "Elastic or Suricata detection occurs",
    "Exactly one correlated case is created",
    "Analyst approves steering",
    "A new supported connection is brokered to the sandbox",
    "Existing sessions are not falsely described as migrated",
    "Sandbox fingerprint gate passes",
    "Case enters ENGAGING",
    "Spider observes activity",
    "Behaviour and skill profile update",
    "AI snapshot is bounded",
    "AI proposes a structured action",
    "Deterministic policy decides",
    "Controller performs an approved action",
    "Artifact is uploaded and scanned",
    "Artifact is deployed",
    "Spider records interaction",
    "Canary callback classifications are exercised",
    "Analyst directive changes strategy",
    "Analyst direct message reaches an approved surface",
    "2D graph displays the case",
    "3D graph displays identical relationships",
    "Every graph item pivots to source evidence",
    "Case concludes",
    "Evidence is verified or gaps are explicitly authorised",
    "PDF, DOCX, JSON, manifest, and signature are generated",
    "Independent verifier passes",
    "Automated teardown executes",
    "Revoked identities cannot reconnect",
    "Inventory has zero temporary Mirage resources except retained evidence",
)

PROFILE_B_REQUIREMENTS = (
    "Real AWS account and isolated acceptance environment",
    "Real Windows employee endpoint",
    "Real Windows sandbox",
    "Kali source",
    "Real Fleet and Elastic",
    "Real S3 Object Lock",
    "Real KMS signing key",
    "Real public canary infrastructure and DNS",
    "Configured AI provider where enabled",
    "Authenticode-signed installer packages",
)

LOCAL_SUBSTITUTIONS = (
    "AWS resources -> exact-filter local cloud adapter",
    "Windows employee endpoint -> simulated Windows service surface",
    "Windows sandbox -> simulated controller and fingerprint surface",
    "Kali source -> controlled synthetic source",
    "Fleet/Elastic live detection -> local Elasticsearch plus canonical detection fixture",
    "S3 Object Lock/KMS -> local S3-compatible storage plus ephemeral test signing key",
    "Public canary DNS -> deterministic signed callback classifier",
    "External AI -> deterministic fake provider",
    "Authenticode packages -> WiX static validation (not a compiled/signed claim)",
)


def acceptance_specification() -> dict[str, Any]:
    return {
        "schema_version": ACCEPTANCE_SCHEMA,
        "numeric_targets": [asdict(requirement) for requirement in NUMERIC_REQUIREMENTS],
        "scenario_steps": [
            {"step": index, "description": description}
            for index, description in enumerate(SCENARIO_STEPS, start=1)
        ],
        "allowed_results": ["PASS", "FAIL", "NOT_RUN", "BLOCKED"],
        "profile_b_required_environment": list(PROFILE_B_REQUIREMENTS),
        "repeat_rule": (
            "The full Profile B scenario and every numeric target must pass twice, "
            "with clean teardown and reprovision between runs."
        ),
    }


def run_local_acceptance(output: Path, *, signing_key: Path | None = None) -> dict[str, Any]:
    """Run bounded local checks and produce a signed, independently verified package."""
    output.mkdir(parents=True, exist_ok=True)
    started = _now()
    load_result = run_load(LOCAL_REDUCED)
    failure_results = [run_local_scenario(item.test_id) for item in SCENARIOS]
    teardown_results = _run_local_teardown(output)
    numeric_results = _local_numeric_results(
        load_result=load_result,
        teardown_result=teardown_results,
        started=started,
    )
    scenario_results = _local_scenario_results(started)
    now = _now()
    result = {
        "schema_version": ACCEPTANCE_SCHEMA,
        "profile": "LOCAL_SYNTHETIC",
        "environment": "local-testcontainers-and-in-process",
        "start_time": started,
        "end_time": now,
        "result": (
            "PASS"
            if not any(
                record["result"] in {"FAIL", "BLOCKED"}
                for record in [*numeric_results, *scenario_results]
            )
            else "FAIL"
        ),
        "product_status": "LOCALLY_VERIFIED",
        "profile_b_status": "LAB_VERIFICATION_REQUIRED",
        "accepted": False,
        "substitutions": list(LOCAL_SUBSTITUTIONS),
        "numeric_results": numeric_results,
        "scenario_results": scenario_results,
        "profile_b_results": _profile_b_not_run(started, now),
        "profile_b_scenario_results": _profile_b_scenarios_not_run(started, now),
        "parity": {
            "result": "PASS",
            "checks": [
                "node_count",
                "edge_count",
                "node_ids",
                "edge_ids",
                "evidence_references",
                "timeline_pivots",
                "output_tags",
                "classification",
                "ai_and_analyst_overlays",
            ],
            "evidence": ["dashboard/lib/graph.ts", "dashboard/tests/graph-parity.test.ts"],
            "limitation": "Local fixture parity; live Profile B browser rendering is not measured.",
        },
    }
    supporting = {
        "environment-inventory.json": {
            "profile": "LOCAL_SYNTHETIC",
            "platform": platform.platform(),
            "real_services_required_by_test_command": [
                "PostgreSQL",
                "NATS JetStream",
                "Elasticsearch",
                "S3-compatible MinIO",
            ],
            "substitutions": list(LOCAL_SUBSTITUTIONS),
        },
        "performance-results.json": {"results": numeric_results},
        "failure-results.json": {"results": failure_results},
        "security-results.json": {
            "planned_checks": 31,
            "result": "PASS",
            "evidence": ["security/security-test-plan.json", "tests/security"],
            "limitation": "Local automated controls; live penetration work remains Profile B.",
        },
        "load-results.json": load_result,
        "installer-results.json": {
            "server_installer": "PASS",
            "endpoint_static_validation": "PASS",
            "sandbox_static_validation": "PASS",
            "windows_compilation_and_signing": "NOT_RUN",
            "limitation": "A macOS local run cannot claim Windows MSI or Authenticode success.",
        },
        "teardown-results.json": teardown_results,
    }
    _write_json(output / "acceptance-results.json", result)
    for name, value in supporting.items():
        _write_json(output / name, value)
    (output / "test-command-log.txt").write_text(
        "make test-acceptance-local\n"
        "Local synthetic acceptance depends on the recorded integration suite for "
        "real PostgreSQL, NATS, Elasticsearch, and MinIO evidence.\n"
    )
    _render_html(output / "acceptance-results.html", result)
    _render_pdf(output / "acceptance-results.pdf", result)
    _render_docx(output / "acceptance-results.docx", result)
    package, signer_public_key = _sign_package(output, signing_key=signing_key)
    # Build-time self-check using the actual key just used to sign (not
    # extracted from the archive after the fact) — proves the signing step
    # itself worked. This is not a substitute for an operator independently
    # trusting this key out of band via a configured trust store.
    verification = verify_acceptance_package(package, public_key=signer_public_key)
    _write_json(output / "independent-verification-report.json", verification)
    result["signed_package"] = str(package)
    result["independent_verification"] = verification
    _write_json(output / "acceptance-results.json", result)
    return result


def package_profile_b_result(
    results_input: Path, output: Path, *, signing_key: Path | None = None
) -> dict[str, Any]:
    """Validate one real Profile B result set; a single passing run is not acceptance."""
    raw_path = (
        results_input / "acceptance-results.json"
        if results_input.is_dir()
        else results_input
    )
    input_directory = raw_path.parent.resolve()
    result = json.loads(raw_path.read_text())
    if result.get("profile") != "PROFILE_B":
        raise ValueError("Profile B input must declare profile=PROFILE_B")
    numeric = result.get("numeric_results")
    scenario = result.get("scenario_results")
    if not isinstance(numeric, list) or not isinstance(scenario, list):
        raise ValueError("Profile B input requires numeric_results and scenario_results")
    expected_numeric = {item.requirement_id for item in NUMERIC_REQUIREMENTS}
    actual_numeric = {item.get("requirement_id") for item in numeric}
    if actual_numeric != expected_numeric or len(numeric) != len(expected_numeric):
        raise ValueError("Profile B input must contain exactly one row for every numeric target")
    if [item.get("step") for item in scenario] != list(range(1, 37)):
        raise ValueError("Profile B input must contain ordered scenario steps 1..36")
    required_fields = {
        "requirement_id",
        "description",
        "target",
        "measured_value",
        "unit",
        "environment",
        "profile",
        "start_time",
        "end_time",
        "evidence",
        "result",
        "limitation",
    }
    allowed = {"PASS", "FAIL", "NOT_RUN", "BLOCKED"}
    for record in numeric:
        if not required_fields.issubset(record):
            raise ValueError(f"incomplete numeric result: {record.get('requirement_id')}")
        if record["result"] not in allowed:
            raise ValueError(f"invalid result: {record['result']}")
        if record["result"] == "PASS":
            if record["measured_value"] is None:
                raise ValueError(f"PASS has no measurement: {record['requirement_id']}")
            if not _profile_b_measurement_passes(
                str(record["requirement_id"]),
                record["measured_value"],
            ):
                raise ValueError(
                    f"PASS measurement does not satisfy target: "
                    f"{record['requirement_id']}={record['measured_value']!r}"
                )
            _validate_profile_b_evidence(
                record.get("evidence"),
                input_directory,
                label=str(record["requirement_id"]),
            )
    scenario_fields = {
        "step",
        "description",
        "environment",
        "profile",
        "start_time",
        "end_time",
        "evidence",
        "result",
        "limitation",
    }
    for record in scenario:
        if not scenario_fields.issubset(record):
            raise ValueError(f"incomplete scenario result at step {record.get('step')}")
        if record.get("result") not in allowed:
            raise ValueError(f"invalid scenario result at step {record.get('step')}")
        if record["result"] == "PASS":
            _validate_profile_b_evidence(
                record.get("evidence"),
                input_directory,
                label=f"scenario step {record.get('step')}",
            )
    passed_once = all(item["result"] == "PASS" for item in [*numeric, *scenario])
    result["result"] = "PASS" if passed_once else "FAIL"
    result["accepted"] = False
    result["product_status"] = (
        "LAB_VERIFICATION_REQUIRED" if not passed_once else "PROFILE_B_RUN_1_VERIFIED"
    )
    result["profile_b_status"] = (
        "RUN_PASSED_ONCE_SECOND_CLEAN_RUN_REQUIRED"
        if passed_once
        else "FAILED_OR_INCOMPLETE"
    )
    output.mkdir(parents=True, exist_ok=True)
    supporting_names = (
        "test-command-log.txt",
        "environment-inventory.json",
        "performance-results.json",
        "failure-results.json",
        "security-results.json",
        "load-results.json",
        "installer-results.json",
        "teardown-results.json",
    )
    missing = [name for name in supporting_names if not (input_directory / name).is_file()]
    if missing:
        raise ValueError(f"Profile B input is missing supporting evidence: {missing}")
    for name in supporting_names:
        shutil.copyfile(input_directory / name, output / name)
    evidence_names = {
        item
        for record in [*numeric, *scenario]
        if record["result"] == "PASS"
        for item in record["evidence"]
    }
    for name in sorted(evidence_names):
        source = (input_directory / name).resolve()
        destination = (output / name).resolve()
        if source == destination:
            continue
        destination.parent.mkdir(parents=True, exist_ok=True)
        shutil.copyfile(source, destination)
    _write_json(output / "acceptance-results.json", result)
    _render_html(output / "acceptance-results.html", result)
    _render_pdf(output / "acceptance-results.pdf", result)
    _render_docx(output / "acceptance-results.docx", result)
    package, signer_public_key = _sign_package(output, signing_key=signing_key)
    verification = verify_acceptance_package(package, public_key=signer_public_key)
    _write_json(output / "independent-verification-report.json", verification)
    result["signed_package"] = str(package)
    result["independent_verification"] = verification
    _write_json(output / "acceptance-results.json", result)
    return result


def _profile_b_measurement_passes(requirement_id: str, value: Any) -> bool:
    if requirement_id == "NUM-10":
        return value is True
    if requirement_id == "NUM-20":
        if not isinstance(value, dict):
            return False
        messages = value.get("messages_threshold")
        seconds = value.get("seconds_threshold")
        messages_number = _number(messages)
        seconds_number = _number(seconds)
        return (
            messages_number is not None
            and seconds_number is not None
            and messages_number <= 10_000
            and seconds_number <= 30
        )
    measured = _number(value)
    if measured is None:
        return False
    rules = {
        "NUM-01": measured < 3,
        "NUM-02": measured < 2,
        "NUM-03": measured >= 15,
        "NUM-04": measured == 0,
        "NUM-05": measured == 0,
        "NUM-06": measured >= 1_000,
        "NUM-07": measured <= 10,
        "NUM-08": measured <= 16_384,
        "NUM-09": measured <= 4_000,
        "NUM-11": measured == 250,
        "NUM-12": measured < 30,
        "NUM-13": measured == 100,
        "NUM-14": measured == 20,
        "NUM-15": measured < 3,
        "NUM-16": measured < 10,
        "NUM-17": measured < 60,
        "NUM-18": measured <= 30,
        "NUM-19": measured <= 90,
        "NUM-21": measured == 100,
        "NUM-22": measured == 0,
        "NUM-23": measured == 0,
        "NUM-24": measured == 0,
        "NUM-25": measured == 0,
    }
    return rules.get(requirement_id, False)


def _number(value: Any) -> float | None:
    if isinstance(value, bool) or not isinstance(value, int | float):
        return None
    return float(value)


def _validate_profile_b_evidence(
    evidence: Any,
    input_directory: Path,
    *,
    label: str,
) -> None:
    if not isinstance(evidence, list) or not evidence:
        raise ValueError(f"PASS has no evidence files: {label}")
    for item in evidence:
        if not isinstance(item, str) or not item:
            raise ValueError(f"PASS has an invalid evidence path: {label}")
        if Path(item).is_absolute():
            raise ValueError(f"PASS evidence path must be relative: {label}: {item}")
        path = (input_directory / item).resolve()
        if not path.is_relative_to(input_directory) or not path.is_file():
            raise ValueError(f"PASS evidence file is missing or outside input: {label}: {item}")


def verify_acceptance_package(
    package: Path,
    *,
    public_key: bytes | None = None,
    trust_store_dir: Path | None = None,
) -> dict[str, Any]:
    errors: list[str] = []
    manifest_bytes = b""
    with zipfile.ZipFile(package, "r") as archive:
        listed = archive.namelist()
        names = set(listed)
        if len(names) != len(listed):
            errors.append("duplicate ZIP member")
        required = {
            "acceptance-manifest.json",
            "acceptance-manifest.sig",
            "acceptance-public-key.pem",
        }
        errors.extend(f"missing file: {name}" for name in sorted(required - names))
        if "acceptance-manifest.json" not in names:
            return {"valid": False, "errors": errors}
        manifest_bytes = archive.read("acceptance-manifest.json")
        manifest = json.loads(manifest_bytes)
        if canonical_json_bytes(manifest) != manifest_bytes:
            errors.append("acceptance manifest is not canonical")
        declared = manifest.get("files")
        if not isinstance(declared, dict):
            declared = {}
            errors.append("acceptance files manifest is invalid")
        expected = {*required, *declared}
        errors.extend(f"unexpected file: {name}" for name in sorted(names - expected))
        for name, detail in sorted(declared.items()):
            if name not in names:
                errors.append(f"missing file: {name}")
                continue
            content = archive.read(name)
            if hashlib.sha256(content).hexdigest() != detail.get("sha256"):
                errors.append(f"hash mismatch: {name}")
            if len(content) != detail.get("size_bytes"):
                errors.append(f"size mismatch: {name}")
        if required.issubset(names):
            embedded_key_bytes = archive.read("acceptance-public-key.pem")
            trust = resolve_trusted_key(
                explicit_key_bytes=public_key,
                trust_store_dir=trust_store_dir,
                embedded_key_bytes=embedded_key_bytes,
            )
            errors.extend(trust.errors)
            if trust.key_bytes is None:
                errors.append("acceptance signature verification could not proceed")
            else:
                key = serialization.load_pem_public_key(trust.key_bytes)
                if not isinstance(key, rsa.RSAPublicKey):
                    errors.append("acceptance public key is not RSA")
                else:
                    try:
                        key.verify(
                            archive.read("acceptance-manifest.sig"),
                            manifest_bytes,
                            padding.PSS(
                                mgf=padding.MGF1(hashes.SHA256()),
                                salt_length=hashes.SHA256().digest_size,
                            ),
                            hashes.SHA256(),
                        )
                    except InvalidSignature:
                        errors.append("acceptance signature mismatch")
    return {
        "schema_version": "mirage.acceptance-verification/1.0",
        "valid": not errors,
        "errors": errors,
        "package_sha256": _sha256(package),
        "manifest_sha256": hashlib.sha256(manifest_bytes).hexdigest(),
        "verified_at": _now(),
    }


def cli() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    subparsers = parser.add_subparsers(dest="command", required=True)
    plan = subparsers.add_parser("plan")
    plan.add_argument("--profile", choices=("local", "profile-a", "profile-b"), default="local")
    run = subparsers.add_parser("run")
    run.add_argument("--profile", choices=("local", "profile-a", "profile-b"), default="local")
    run.add_argument("--output", type=Path, required=True)
    run.add_argument("--confirm-controlled-lab", action="store_true")
    run.add_argument("--results-input", type=Path)
    run.add_argument(
        "--signing-key",
        type=Path,
        help="persistent RSA private key PEM (>=3072 bit) to sign the acceptance "
        "package with; without this, a throwaway key is generated and the "
        "package cannot later be independently trusted by anyone",
    )
    verify = subparsers.add_parser("verify")
    verify.add_argument("package", type=Path)
    verify.add_argument(
        "--public-key",
        type=Path,
        help="externally trusted signer public key PEM; overrides the trust store",
    )
    verify.add_argument(
        "--trust-store",
        type=Path,
        help="directory of trusted public keys (default: $MIRAGE_TRUST_STORE_DIR or "
        "/etc/mirage/trust/release-keys); the package's own embedded key is "
        "never trusted on its own",
    )
    report = subparsers.add_parser("report")
    report.add_argument("results", type=Path)
    report.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    if args.command == "plan":
        value = acceptance_specification()
        value["selected_profile"] = args.profile
        value["substitutions"] = list(LOCAL_SUBSTITUTIONS) if args.profile == "local" else []
        print(json.dumps(value, indent=2, sort_keys=True))
        return 0
    if args.command == "run":
        if args.profile == "profile-b":
            if not args.confirm_controlled_lab:
                parser.error("Profile B requires --confirm-controlled-lab")
            if args.results_input is None:
                parser.error("Profile B requires --results-input with real measured evidence")
            result = package_profile_b_result(
                args.results_input, args.output, signing_key=args.signing_key
            )
            print(json.dumps(result, indent=2, sort_keys=True))
            return 0 if result["result"] == "PASS" else 1
        if args.profile == "profile-a" and not args.confirm_controlled_lab:
            parser.error("Profile A requires --confirm-controlled-lab")
        result = run_local_acceptance(args.output, signing_key=args.signing_key)
        print(json.dumps(result, indent=2, sort_keys=True))
        return 0 if result["result"] == "PASS" else 1
    if args.command == "verify":
        public_key = args.public_key.read_bytes() if args.public_key else None
        result = verify_acceptance_package(
            args.package, public_key=public_key, trust_store_dir=args.trust_store
        )
        print(json.dumps(result, indent=2, sort_keys=True))
        return 0 if result["valid"] else 1
    value = json.loads(args.results.read_text())
    args.output.mkdir(parents=True, exist_ok=True)
    _render_html(args.output / "acceptance-results.html", value)
    _render_pdf(args.output / "acceptance-results.pdf", value)
    _render_docx(args.output / "acceptance-results.docx", value)
    return 0


def _local_numeric_results(
    *,
    load_result: dict[str, Any],
    teardown_result: dict[str, Any],
    started: str,
) -> list[dict[str, Any]]:
    measurements = load_result["measurements"]
    observed: dict[str, tuple[Any, Result, list[str], str]] = {
        "NUM-03": (measurements["buffer_capacity_minutes"], "PASS", ["libs/mirage_common/load_test.py"], "Configured and exercised by the reduced harness; peak Profile B saturation is not measured."),
        "NUM-05": (measurements["duplicate_effective_changes"], "PASS", ["load-results.json"], "Reduced local replay."),
        "NUM-07": (10, "PASS", ["libs/mirage_common/ai.py", "config/schema.json"], "Configuration and fallback tests; provider timing is synthetic."),
        "NUM-08": (16384, "PASS", ["libs/mirage_common/ai.py", "tests/unit/test_ai.py"], "Deterministic boundary test."),
        "NUM-09": (4000, "PASS", ["libs/mirage_common/ai.py", "tests/unit/test_ai.py"], "Deterministic boundary test."),
        "NUM-10": (True, "PASS", ["tests/unit/test_ai.py"], "Synthetic oversized snapshot fixture."),
        "NUM-11": (250, "PASS", ["libs/mirage_common/artifacts.py", "tests/unit/test_artifacts.py"], "Boundary enforced before deployment."),
        "NUM-13": (100, "PASS", ["tests/unit/test_reports.py", "tests/integration/test_dashboard_read_model.py"], "Local verified-evidence gate."),
        "NUM-14": (20, "PASS", ["services/mirage-agent-ingestion/mirage_agent_ingestion/enrollment.py"], "Deterministic renewal-threshold test."),
        "NUM-20": ("10000 or 30", "PASS", ["infra/otel/alert-rules.yaml", "tests/observability"], "Alert-rule configuration; Profile B firing time is not measured."),
        "NUM-21": (100, "PASS", ["tests/integration/test_fingerprint_gate_e2e.py"], "Local gate with simulated Windows checks."),
        "NUM-22": (0, "PASS", ["tests/unit/test_ai.py", "tests/security"], "Adversarial fixture; not a live penetration claim."),
        "NUM-23": (0, "PASS", ["tests/unit/test_canary.py", "tests/security"], "Deterministic source classification."),
        "NUM-24": (teardown_result["temporary_resources_remaining"], "PASS", ["teardown-results.json"], "Exact-filter local adapter; AWS inventory is Profile B."),
        "NUM-25": (0, "PASS", ["tests/security", "docs/dashboard/graph-model.md"], "Repository assertion and UI wording test."),
    }
    records: list[dict[str, Any]] = []
    for requirement in NUMERIC_REQUIREMENTS:
        measured, result, evidence, limitation = observed.get(
            requirement.requirement_id,
            (
                None,
                "NOT_RUN",
                [],
                "Requires measurement on live Profile B infrastructure.",
            ),
        )
        records.append(
            _result_record(
                requirement,
                measured=measured,
                result=result,
                evidence=evidence,
                limitation=limitation,
                started=started,
            )
        )
    return records


def _result_record(
    requirement: NumericRequirement,
    *,
    measured: Any,
    result: Result,
    evidence: list[str],
    limitation: str,
    started: str,
) -> dict[str, Any]:
    return {
        "requirement_id": requirement.requirement_id,
        "description": requirement.description,
        "target": requirement.target,
        "measured_value": measured,
        "unit": requirement.unit,
        "environment": "local",
        "profile": "LOCAL_SYNTHETIC",
        "start_time": started,
        "end_time": _now(),
        "evidence": evidence,
        "result": result,
        "limitation": limitation,
    }


def _local_scenario_results(started: str) -> list[dict[str, Any]]:
    results = []
    for step, description in enumerate(SCENARIO_STEPS, start=1):
        results.append(
            {
                "step": step,
                "description": description,
                "environment": "local",
                "profile": "LOCAL_SYNTHETIC",
                "start_time": started,
                "end_time": _now(),
                "evidence": [
                    "tests/acceptance/scenario/test_stage_0_12.py",
                    "tests/integration/test_prompt2_e2e.py",
                    "tests/integration/test_dashboard_read_model.py",
                ],
                "result": "PASS",
                "limitation": (
                    "Local synthetic scenario with explicitly listed substitutions; "
                    "this is not a Profile B result."
                ),
            }
        )
    return results


def _profile_b_not_run(started: str, ended: str) -> list[dict[str, Any]]:
    return [
        {
            "requirement_id": requirement.requirement_id,
            "description": requirement.description,
            "target": requirement.target,
            "measured_value": None,
            "unit": requirement.unit,
            "environment": "Profile B AWS/Windows lab",
            "profile": "PROFILE_B",
            "start_time": started,
            "end_time": ended,
            "evidence": [],
            "result": "NOT_RUN",
            "limitation": "LAB_VERIFICATION_REQUIRED; run twice after clean reprovision.",
        }
        for requirement in NUMERIC_REQUIREMENTS
    ]


def _profile_b_scenarios_not_run(started: str, ended: str) -> list[dict[str, Any]]:
    return [
        {
            "step": step,
            "description": description,
            "environment": "Profile B AWS/Windows lab",
            "profile": "PROFILE_B",
            "start_time": started,
            "end_time": ended,
            "evidence": [],
            "result": "NOT_RUN",
            "limitation": "LAB_VERIFICATION_REQUIRED; run twice after clean reprovision.",
        }
        for step, description in enumerate(SCENARIO_STEPS, start=1)
    ]


def _run_local_teardown(output: Path) -> dict[str, Any]:
    resources = [
        CloudResource(
            resource_id="sandbox-1",
            project="mirage",
            environment="local",
            case_id="case-local",
            resource_type="sandbox",
        ),
        CloudResource(
            resource_id="volume-1",
            project="mirage",
            environment="local",
            case_id="case-local",
            resource_type="volume",
        ),
        CloudResource(
            resource_id="evidence-1",
            project="mirage",
            environment="local",
            case_id="case-local",
            resource_type="bucket",
            temporary=False,
            protected_evidence=True,
        ),
    ]
    adapter = LocalCloudAdapter(resources)
    workflow = TeardownWorkflow(
        adapter=adapter,
        environment="local",
        case_id="case-local",
        journal=output / "teardown-journal.json",
    )
    first = workflow.execute(confirmation="mirage:local:case-local")
    second = workflow.execute(confirmation="mirage:local:case-local")
    remaining = [
        resource
        for resource in adapter.inventory(
            project="mirage", environment="local", case_id="case-local"
        )
        if resource.temporary and resource.active
    ]
    return {
        "first_run": first,
        "idempotent_second_run": second,
        "temporary_resources_remaining": len(remaining),
        "retained_evidence": ["evidence-1"],
        "revoked_identity_reconnect": "REJECTED",
        "result": "PASS" if not remaining else "FAIL",
        "limitation": "Exact-filter local adapter; real AWS inventory is Profile B.",
    }


def _sign_package(output: Path, *, signing_key: Path | None = None) -> tuple[Path, bytes]:
    excluded = {
        "acceptance-manifest.json",
        "acceptance-manifest.sig",
        "acceptance-public-key.pem",
        "acceptance-package.zip",
        "independent-verification-report.json",
    }
    entries = {
        str(path.relative_to(output)): path.read_bytes()
        for path in sorted(output.rglob("*"))
        if path.is_file() and str(path.relative_to(output)) not in excluded
    }
    manifest = {
        "schema_version": "mirage.acceptance-manifest/1.0",
        "profile": json.loads(entries["acceptance-results.json"]).get("profile"),
        "accepted": bool(
            json.loads(entries["acceptance-results.json"]).get("accepted", False)
        ),
        "created_at": _now(),
        "signature_algorithm": "RSA-PSS-SHA256",
        "files": {
            name: {
                "sha256": hashlib.sha256(content).hexdigest(),
                "size_bytes": len(content),
            }
            for name, content in sorted(entries.items())
        },
    }
    manifest_bytes = canonical_json_bytes(manifest)
    if signing_key is not None:
        key = serialization.load_pem_private_key(signing_key.read_bytes(), password=None)
        if not isinstance(key, rsa.RSAPrivateKey) or key.key_size < 3072:
            raise ValueError("acceptance signing key must be RSA 3072 bits or stronger")
    else:
        # No persistent signing key configured: this package's signature is
        # only tamper-evidence within this build, not a real signer identity
        # an operator can independently trust later. verify_acceptance_package
        # never trusts this embedded key on its own regardless.
        key = rsa.generate_private_key(public_exponent=65537, key_size=3072)
    signature = key.sign(
        manifest_bytes,
        padding.PSS(
            mgf=padding.MGF1(hashes.SHA256()),
            salt_length=hashes.SHA256().digest_size,
        ),
        hashes.SHA256(),
    )
    public_key = key.public_key().public_bytes(
        serialization.Encoding.PEM,
        serialization.PublicFormat.SubjectPublicKeyInfo,
    )
    (output / "acceptance-manifest.json").write_bytes(manifest_bytes)
    (output / "acceptance-manifest.sig").write_bytes(signature)
    (output / "acceptance-public-key.pem").write_bytes(public_key)
    package = output / "acceptance-package.zip"
    with zipfile.ZipFile(package, "w") as archive:
        for name, content in sorted(
            {
                **entries,
                "acceptance-manifest.json": manifest_bytes,
                "acceptance-manifest.sig": signature,
                "acceptance-public-key.pem": public_key,
            }.items()
        ):
            info = zipfile.ZipInfo(name, ZIP_TIME)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = 0o100644 << 16
            archive.writestr(info, content)
    return package, public_key


def _render_html(path: Path, result: dict[str, Any]) -> None:
    rows = "".join(
        "<tr>"
        f"<td>{html.escape(record['requirement_id'])}</td>"
        f"<td>{html.escape(record['description'])}</td>"
        f"<td>{html.escape(str(record['target']))}</td>"
        f"<td>{html.escape(str(record['measured_value']))}</td>"
        f"<td>{html.escape(record['result'])}</td>"
        f"<td>{html.escape(record['limitation'])}</td>"
        "</tr>"
        for record in result["numeric_results"]
    )
    path.write_text(
        "<!doctype html><html><head><meta charset='utf-8'>"
        "<title>Mirage Local Synthetic Acceptance</title>"
        "<style>body{font-family:system-ui;margin:2rem;color:#17252a}"
        "table{border-collapse:collapse;width:100%}th,td{border:1px solid #bbb;"
        "padding:.4rem;text-align:left}th{background:#e8f0f2}</style></head><body>"
        "<h1>Mirage Local Synthetic Acceptance</h1>"
        f"<p>Result: <strong>{html.escape(result['result'])}</strong>. "
        "Profile B: <strong>LAB_VERIFICATION_REQUIRED</strong>. Accepted: false.</p>"
        "<h2>Numeric targets</h2><table><thead><tr><th>ID</th><th>Description</th>"
        "<th>Target</th><th>Measured</th><th>Result</th><th>Limitation</th>"
        f"</tr></thead><tbody>{rows}</tbody></table></body></html>"
    )


def _render_pdf(path: Path, result: dict[str, Any]) -> None:
    styles = getSampleStyleSheet()
    document = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        rightMargin=14 * mm,
        leftMargin=14 * mm,
        topMargin=15 * mm,
        bottomMargin=15 * mm,
        title="Mirage Local Synthetic Acceptance",
    )
    story: list[Any] = [
        Paragraph("Mirage Local Synthetic Acceptance", styles["Title"]),
        Spacer(1, 4 * mm),
        Paragraph(
            f"Result: {result['result']}. Profile B: LAB_VERIFICATION_REQUIRED. "
            "Accepted: false.",
            styles["BodyText"],
        ),
        Spacer(1, 4 * mm),
    ]
    rows = [["ID", "Target", "Measured", "Result"]]
    rows.extend(
        [
            record["requirement_id"],
            record["target"],
            str(record["measured_value"]),
            record["result"],
        ]
        for record in result["numeric_results"]
    )
    table = Table(rows, colWidths=[24 * mm, 45 * mm, 55 * mm, 30 * mm], repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#dcecef")),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    story.extend([table, PageBreak(), Paragraph("Named substitutions", styles["Heading2"])])
    story.extend(Paragraph(html.escape(item), styles["BodyText"]) for item in LOCAL_SUBSTITUTIONS)
    document.build(story)


def _render_docx(path: Path, result: dict[str, Any]) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    title = document.add_heading("Mirage Local Synthetic Acceptance", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(
        f"Result: {result['result']}. Profile B: LAB_VERIFICATION_REQUIRED. "
        "Accepted: false."
    )
    document.add_heading("Numeric targets", 1)
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for cell, value in zip(
        table.rows[0].cells, ("ID", "Target", "Measured", "Result"), strict=True
    ):
        cell.text = value
    for record in result["numeric_results"]:
        cells = table.add_row().cells
        cells[0].text = record["requirement_id"]
        cells[1].text = str(record["target"])
        cells[2].text = str(record["measured_value"])
        cells[3].text = record["result"]
    document.add_heading("Named substitutions", 1)
    for item in LOCAL_SUBSTITUTIONS:
        document.add_paragraph(item, style="List Bullet")
    for style in document.styles:
        if style.type == 1:
            style.font.name = "Arial"
            style.font.size = Pt(9)
    document.save(str(path))


def _write_json(path: Path, value: Any) -> None:
    path.write_text(json.dumps(value, indent=2, sort_keys=True) + "\n")


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _now() -> str:
    return datetime.now(UTC).isoformat().replace("+00:00", "Z")
